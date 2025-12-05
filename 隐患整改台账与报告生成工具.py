import os
import zipfile
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XLImage
from PIL import Image as PILImage
import io
import warnings
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import pandas as pd

# 启用 HEIC 支持
try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
except ImportError:
    pass

warnings.filterwarnings("ignore", category=UserWarning, module='PIL')

# 图片目标尺寸：3.5cm x 5cm → 像素（96 DPI）
CM_TO_PIXEL = 96 / 2.54
IMG_WIDTH_PX = int(5 * CM_TO_PIXEL)    # ≈ 189 (5cm宽度)
IMG_HEIGHT_PX = int(3.5 * CM_TO_PIXEL) # ≈ 132 (3.5cm高度)

# openpyxl 行列尺寸参考值
COL_WIDTH_FOR_IMG = 35   # E 和 M 列宽度
ROW_HEIGHT_FOR_IMG = 100 # 插入图片的行高

def embed_images_to_excel(excel_path, zip_path, output_path):
    wb = load_workbook(excel_path)
    ws = wb.active

    if ws.cell(row=1, column=1).value != "隐患编号":
        raise ValueError("Excel 第一列标题必须是“隐患编号”")

    folder_images = extract_zip_to_dict(zip_path)
    risk_photos = folder_images["隐患照片"]
    close_loop_photos = folder_images["闭环照片"]

    # 初始化错误列表
    errors = []

    # 先设置列宽（E列=隐患照片, M列=闭环照片）
    ws.column_dimensions['E'].width = COL_WIDTH_FOR_IMG
    ws.column_dimensions['M'].width = COL_WIDTH_FOR_IMG

    row_idx = 2
    while True:
        cell_value = ws.cell(row=row_idx, column=1).value
        if cell_value is None:
            break
        key = str(cell_value).strip()

        inserted = False

        # 隐患照片 → E列
        if key in risk_photos:
            try:
                img_data = risk_photos[key]
                pil_img = PILImage.open(io.BytesIO(img_data))
                pil_img = pil_img.resize((IMG_WIDTH_PX, IMG_HEIGHT_PX), PILImage.LANCZOS)
                if pil_img.mode in ("RGBA", "P"):
                    pil_img = pil_img.convert("RGB")
                img_buffer = io.BytesIO()
                pil_img.save(img_buffer, format='JPEG')
                img_buffer.seek(0)
                xl_img = XLImage(img_buffer)
                xl_img.width = IMG_WIDTH_PX
                xl_img.height = IMG_HEIGHT_PX
                ws.add_image(xl_img, f"E{row_idx}")
                inserted = True
            except Exception as e:
                errors.append(f"第 {row_idx} 行（隐患编号 {key}）隐患照片插入失败: {str(e)}")

        # 闭环照片 → M列
        if key in close_loop_photos:
            try:
                img_data = close_loop_photos[key]
                pil_img = PILImage.open(io.BytesIO(img_data))
                pil_img = pil_img.resize((IMG_WIDTH_PX, IMG_HEIGHT_PX), PILImage.LANCZOS)
                if pil_img.mode in ("RGBA", "P"):
                    pil_img = pil_img.convert("RGB")
                img_buffer = io.BytesIO()
                pil_img.save(img_buffer, format='JPEG')
                img_buffer.seek(0)
                xl_img = XLImage(img_buffer)
                xl_img.width = IMG_WIDTH_PX
                xl_img.height = IMG_HEIGHT_PX
                ws.add_image(xl_img, f"M{row_idx}")
                inserted = True
            except Exception as e:
                errors.append(f"第 {row_idx} 行（隐患编号 {key}）闭环照片插入失败: {str(e)}")

        # 如果该行插入了图片，设置行高
        if inserted:
            ws.row_dimensions[row_idx].height = ROW_HEIGHT_FOR_IMG

        row_idx += 1

    wb.save(output_path)

    return errors

def generate_check_report(excel_path, doc_template_path, zip_path, output_path):
    # 读取Excel文件
    df = pd.read_excel(excel_path)
    
    # 检查关键列是否存在
    required_cols = ['隐患编号', '异常类别', '隐患级别', '异常事项', '班组', '整改人', '发现时间', '要求闭环时间']
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        raise ValueError(f"Excel文件中缺少以下列: {missing_cols}")
    
    # 将日期列转换为指定格式
    df = df.copy()
    if '发现时间' in df.columns:
        df['发现时间'] = pd.to_datetime(df['发现时间'], errors='coerce').dt.strftime('%Y-%m-%d')
    if '要求闭环时间' in df.columns:
        df['要求闭环时间'] = pd.to_datetime(df['要求闭环时间'], errors='coerce').dt.strftime('%Y-%m-%d')
    
    # 加载Word模板
    doc = Document(doc_template_path)
    
    # 提取ZIP中的图片
    folder_images = extract_zip_to_dict(zip_path)
    
    # 通过标题段落查找表格
    env_table = find_table_by_title(doc, "二、环境保护")
    general_table = find_table_by_title(doc, "一、本期存在主要问题")
    major_table = find_table_by_title(doc, "三、重大事故隐患检查情况")
    
    # 初始化计数器
    env_counter = 0
    general_counter = 0
    major_counter = 0
    
    # 处理每一行数据
    for index, row in df.iterrows():
        # 检查异常类别是否为"环境保护"
        if str(row.get('异常类别', '')) == '环境保护':
            if env_table:
                env_counter += 1
                add_row_to_table_with_images_from_zip(env_table, row, env_counter, folder_images, "隐患照片")
        
        # 检查隐患级别是否为"一般隐患"且异常类别不为"环境保护"
        if str(row.get('隐患级别', '')) == '一般隐患' and str(row.get('异常类别', '')) != '环境保护':
            if general_table:
                general_counter += 1
                add_row_to_table_with_images_from_zip(general_table, row, general_counter, folder_images, "隐患照片")
        
        # 检查隐患级别是否为"重大隐患"且异常类别不为"环境保护"
        if str(row.get('隐患级别', '')) == '重大隐患' and str(row.get('异常类别', '')) != '环境保护':
            if major_table:
                major_counter += 1
                add_row_to_table_with_images_from_zip(major_table, row, major_counter, folder_images, "隐患照片")
    
    # 如果环境保护表格为空，添加一行说明
    if env_table and env_counter == 0:
        new_row = env_table.add_row()
        cells = new_row.cells
        if len(cells) >= 8:  # 确保表格有足够的列
            cells[0].text = "1"
            cells[1].text = "本次检查未发现公司存在环境保护相关问题"
            # 填充其余列
            for i in range(2, len(cells)):
                cells[i].text = "/"
    
    # 如果重大事故隐患检查情况表格为空，添加一行说明
    if major_table and major_counter == 0:
        new_row = major_table.add_row()
        cells = new_row.cells
        if len(cells) >= 8:  # 确保表格有足够的列
            cells[0].text = "1"
            cells[1].text = "根据《重大事故隐患清单》逐一排查，发现公司未存在重大事故隐患。"
            # 填充其余列
            for i in range(2, len(cells)):
                cells[i].text = "/"
    
    # 为所有表格设置格式
    if env_table:
        apply_table_formatting(env_table)
    if general_table:
        apply_table_formatting(general_table)
    if major_table:
        apply_table_formatting(major_table)
    
    # 保存文档
    doc.save(output_path)

def generate_closure_report(excel_path, doc_template_path, zip_path, output_path):
    # 读取Excel文件
    df = pd.read_excel(excel_path)
    
    # 检查关键列是否存在
    required_cols = ['隐患编号', '异常类别', '隐患级别', '异常事项', '班组', '整改人', '发现时间', '要求闭环时间']
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        raise ValueError(f"Excel文件中缺少以下列: {missing_cols}")
    
    # 将日期列转换为指定格式
    df = df.copy()
    if '发现时间' in df.columns:
        df['发现时间'] = pd.to_datetime(df['发现时间'], errors='coerce').dt.strftime('%Y-%m-%d')
    if '要求闭环时间' in df.columns:
        df['要求闭环时间'] = pd.to_datetime(df['要求闭环时间'], errors='coerce').dt.strftime('%Y-%m-%d')
    
    # 加载Word模板
    doc = Document(doc_template_path)
    
    # 提取ZIP中的图片
    folder_images = extract_zip_to_dict(zip_path)
    
    # 通过标题段落查找表格
    env_table = find_table_by_title(doc, "二、环境保护")
    general_table = find_table_by_title(doc, "一、本期存在主要问题")
    major_table = find_table_by_title(doc, "三、重大事故隐患检查情况")
    
    # 初始化计数器
    env_counter = 0
    general_counter = 0
    major_counter = 0
    
    # 处理每一行数据
    for index, row in df.iterrows():
        # 检查异常类别是否为"环境保护"
        if str(row.get('异常类别', '')) == '环境保护':
            if env_table:
                env_counter += 1
                add_row_to_table_with_images_from_both_zip(env_table, row, env_counter, folder_images)
        
        # 检查隐患级别是否为"一般隐患"且异常类别不为"环境保护"
        if str(row.get('隐患级别', '')) == '一般隐患' and str(row.get('异常类别', '')) != '环境保护':
            if general_table:
                general_counter += 1
                add_row_to_table_with_images_from_both_zip(general_table, row, general_counter, folder_images)
        
        # 检查隐患级别是否为"重大隐患"且异常类别不为"环境保护"
        if str(row.get('隐患级别', '')) == '重大隐患' and str(row.get('异常类别', '')) != '环境保护':
            if major_table:
                major_counter += 1
                add_row_to_table_with_images_from_both_zip(major_table, row, major_counter, folder_images)
    
    # 如果环境保护表格为空，添加一行说明
    if env_table and env_counter == 0:
        new_row = env_table.add_row()
        cells = new_row.cells
        if len(cells) >= 9:  # 确保表格有足够的列
            cells[0].text = "1"
            cells[1].text = "本次检查未发现公司存在环境保护相关问题"
            # 填充其余列
            for i in range(2, len(cells)):
                cells[i].text = "/"
    
    # 如果重大事故隐患检查情况表格为空，添加一行说明
    if major_table and major_counter == 0:
        new_row = major_table.add_row()
        cells = new_row.cells
        if len(cells) >= 9:  # 确保表格有足够的列
            cells[0].text = "1"
            cells[1].text = "根据《重大事故隐患清单》逐一排查，发现公司未存在重大事故隐患。"
            # 填充其余列
            for i in range(2, len(cells)):
                cells[i].text = "/"
    
    # 为所有表格设置格式
    if env_table:
        apply_table_formatting(env_table)
    if general_table:
        apply_table_formatting(general_table)
    if major_table:
        apply_table_formatting(major_table)
    
    # 保存文档
    doc.save(output_path)

def find_table_by_title(doc, title_text):
    """
    通过标题段落查找表格
    """
    # 遍历所有段落，查找标题
    for paragraph in doc.paragraphs:
        if title_text in paragraph.text:
            # 从该段落之后开始查找表格
            paragraph_element = paragraph._element
            parent = paragraph_element.getparent()
            
            # 获取段落在父元素中的位置
            paragraph_index = parent.index(paragraph_element)
            
            # 从段落之后开始查找表格元素
            for j in range(paragraph_index + 1, len(parent)):
                element = parent[j]
                # 检查是否是表格元素
                if element.tag.endswith('tbl'):
                    # 现在需要找到对应的Table对象
                    for table in doc.tables:
                        if table._element == element:
                            return table
            break
            
    return None

def find_column_index(table, header_text):
    """
    在表格中查找指定列标题的索引
    """
    if table.rows:
        header_row = table.rows[0]
        for i, cell in enumerate(header_row.cells):
            if header_text in cell.text:
                return i
    return -1

def add_row_to_table_with_images_from_zip(table, row, serial_number, folder_images, photo_folder):
    """向表格添加一行数据并插入图片（检查报告）"""
    new_row = table.add_row()
    cells = new_row.cells
    
    # 第一列填充序号
    cells[0].text = str(serial_number)
    
    # 填充数据 - 根据实际列名映射
    cells[1].text = str(row.get('异常事项', '')) if pd.notna(row.get('异常事项', '')) else ""
    cells[2].text = str(row.get('异常类别', '')) if pd.notna(row.get('异常类别', '')) else ""
    cells[3].text = str(row.get('班组', '')) if pd.notna(row.get('班组', '')) else ""
    cells[4].text = str(row.get('整改人', '')) if pd.notna(row.get('整改人', '')) else ""
    cells[5].text = str(row.get('发现时间', '')) if pd.notna(row.get('发现时间', '')) else ""
    cells[6].text = str(row.get('要求闭环时间', '')) if pd.notna(row.get('要求闭环时间', '')) else ""
    
    # 查找隐患照片列的索引
    hazard_photo_col = find_column_index(table, "隐患照片")
    
    # 插入隐患照片
    if hazard_photo_col != -1 and str(row.get('隐患编号', '')) in folder_images[photo_folder]:
        try:
            # 清空单元格文本
            cells[hazard_photo_col].text = ""
            # 获取图片数据
            hazard_id = str(row.get('隐患编号', ''))
            img_data = folder_images[photo_folder][hazard_id]
            pil_img = PILImage.open(io.BytesIO(img_data))
            pil_img = pil_img.resize((IMG_WIDTH_PX, IMG_HEIGHT_PX), PILImage.LANCZOS)
            if pil_img.mode in ("RGBA", "P"):
                pil_img = pil_img.convert("RGB")
            img_buffer = io.BytesIO()
            pil_img.save(img_buffer, format='JPEG')
            img_buffer.seek(0)
            # 插入图片
            run = cells[hazard_photo_col].paragraphs[0].add_run()
            # 图片尺寸：3.5cm x 5cm
            run.add_picture(img_buffer, width=Inches(1.97), height=Inches(1.38))  # 5cm=1.97英寸, 3.5cm=1.38英寸
        except Exception as e:
            cells[hazard_photo_col].text = ""
    elif hazard_photo_col != -1:
        cells[hazard_photo_col].text = ""

def add_row_to_table_with_images_from_both_zip(table, row, serial_number, folder_images):
    """向表格添加一行数据并插入两种图片（闭环报告）"""
    new_row = table.add_row()
    cells = new_row.cells
    
    # 第一列填充序号
    cells[0].text = str(serial_number)
    
    # 填充数据 - 根据实际列名映射
    cells[1].text = str(row.get('异常事项', '')) if pd.notna(row.get('异常事项', '')) else ""
    cells[2].text = str(row.get('异常类别', '')) if pd.notna(row.get('异常类别', '')) else ""
    cells[3].text = str(row.get('班组', '')) if pd.notna(row.get('班组', '')) else ""
    cells[4].text = str(row.get('整改人', '')) if pd.notna(row.get('整改人', '')) else ""
    cells[5].text = str(row.get('发现时间', '')) if pd.notna(row.get('发现时间', '')) else ""
    cells[6].text = str(row.get('要求闭环时间', '')) if pd.notna(row.get('要求闭环时间', '')) else ""
    
    # 查找隐患照片列和闭环照片列的索引
    hazard_photo_col = find_column_index(table, "隐患照片")
    close_photo_col = find_column_index(table, "闭环照片")
    
    hazard_id = str(row.get('隐患编号', ''))
    
    # 插入隐患照片
    if hazard_photo_col != -1 and hazard_id in folder_images["隐患照片"]:
        try:
            # 清空单元格文本
            cells[hazard_photo_col].text = ""
            # 获取图片数据
            img_data = folder_images["隐患照片"][hazard_id]
            pil_img = PILImage.open(io.BytesIO(img_data))
            pil_img = pil_img.resize((IMG_WIDTH_PX, IMG_HEIGHT_PX), PILImage.LANCZOS)
            if pil_img.mode in ("RGBA", "P"):
                pil_img = pil_img.convert("RGB")
            img_buffer = io.BytesIO()
            pil_img.save(img_buffer, format='JPEG')
            img_buffer.seek(0)
            # 插入图片
            run = cells[hazard_photo_col].paragraphs[0].add_run()
            # 图片尺寸：3.5cm x 5cm
            run.add_picture(img_buffer, width=Inches(1.97), height=Inches(1.38))  # 5cm=1.97英寸, 3.5cm=1.38英寸
        except Exception as e:
            cells[hazard_photo_col].text = ""
    elif hazard_photo_col != -1:
        cells[hazard_photo_col].text = ""
    
    # 插入闭环照片
    if close_photo_col != -1 and hazard_id in folder_images["闭环照片"]:
        try:
            # 清空单元格文本
            cells[close_photo_col].text = ""
            # 获取图片数据
            img_data = folder_images["闭环照片"][hazard_id]
            pil_img = PILImage.open(io.BytesIO(img_data))
            pil_img = pil_img.resize((IMG_WIDTH_PX, IMG_HEIGHT_PX), PILImage.LANCZOS)
            if pil_img.mode in ("RGBA", "P"):
                pil_img = pil_img.convert("RGB")
            img_buffer = io.BytesIO()
            pil_img.save(img_buffer, format='JPEG')
            img_buffer.seek(0)
            # 插入图片
            run = cells[close_photo_col].paragraphs[0].add_run()
            # 图片尺寸：3.5cm x 5cm
            run.add_picture(img_buffer, width=Inches(1.97), height=Inches(1.38))  # 5cm=1.97英寸, 3.5cm=1.38英寸
        except Exception as e:
            cells[close_photo_col].text = ""
    elif close_photo_col != -1:
        cells[close_photo_col].text = ""

def apply_table_formatting(table):
    """为表格应用格式"""
    for row in table.rows:
        for cell in row.cells:
            # 设置单元格内的所有段落格式
            for paragraph in cell.paragraphs:
                # 设置段落居中
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    # 设置字体为仿宋，四号字
                    run.font.name = '仿宋'
                    run.font.size = Pt(14)  # 四号字约等于14磅
                    # 设置加粗
                    run.font.bold = True
                    # 设置中文字体
                    if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                        if run._element.rPr.rFonts is not None:
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def extract_zip_to_dict(zip_path):
    required_folders = {"隐患照片", "闭环照片"}
    folder_images = {"隐患照片": {}, "闭环照片": {}}

    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        all_files = zip_ref.namelist()
        top_dirs = set(f.split('/')[0] for f in all_files if '/' in f)

        # 初始化所需文件夹的字典，即使它们不存在于ZIP文件中
        for folder in required_folders:
            if folder not in top_dirs:
                print(f"警告: ZIP 中缺少 '{folder}' 文件夹，将跳过相关处理。")
                continue  # 跳过不存在的文件夹

            for f in all_files:
                if f.startswith(folder + '/') and not f.endswith('/'):
                    filename = os.path.basename(f)
                    key = os.path.splitext(filename)[0]
                    with zip_ref.open(f) as img_file:
                        folder_images[folder][key] = img_file.read()

    return folder_images

# GUI 应用
class App:
    def __init__(self, root):
        self.root = root
        self.root.title("隐患整改台账与报告生成工具")
        
        # 存储文件路径
        self.excel_path = ""
        self.check_report_template_path = ""
        self.closure_report_template_path = ""
        self.zip_path = ""
        
        self.setup_ui()

    def setup_ui(self):
        frame = ttk.Frame(self.root, padding="20")
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # 隐患整改通知单选择
        ttk.Button(frame, text="选择隐患整改通知单", command=self.select_excel).grid(row=0, column=0, pady=5, sticky=tk.W)
        self.excel_label = ttk.Label(frame, text="未选择")
        self.excel_label.grid(row=0, column=1, padx=10, sticky=tk.W)

        # 检查报告模板选择
        ttk.Button(frame, text="选择检查报告模板", command=self.select_check_report_template).grid(row=1, column=0, pady=5, sticky=tk.W)
        self.check_report_template_label = ttk.Label(frame, text="未选择")
        self.check_report_template_label.grid(row=1, column=1, padx=10, sticky=tk.W)

        # 闭环报告模板选择
        ttk.Button(frame, text="选择闭环报告模板", command=self.select_closure_report_template).grid(row=2, column=0, pady=5, sticky=tk.W)
        self.closure_report_template_label = ttk.Label(frame, text="未选择")
        self.closure_report_template_label.grid(row=2, column=1, padx=10, sticky=tk.W)

        # 图片压缩包选择
        ttk.Button(frame, text="选择图片压缩包", command=self.select_zip).grid(row=3, column=0, pady=5, sticky=tk.W)
        self.zip_label = ttk.Label(frame, text="未选择")
        self.zip_label.grid(row=3, column=1, padx=10, sticky=tk.W)

        # 按钮区域
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=4, column=0, columnspan=2, pady=20)

        ttk.Button(button_frame, text="隐患整改台账生成", command=self.generate_excel_report).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="检查报告生成", command=self.generate_check_report).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="闭环报告生成", command=self.generate_closure_report).pack(side=tk.LEFT, padx=5)

    def select_excel(self):
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if path:
            self.excel_path = path
            self.excel_label.config(text=os.path.basename(path))

    def select_check_report_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if path:
            self.check_report_template_path = path
            self.check_report_template_label.config(text=os.path.basename(path))

    def select_closure_report_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if path:
            self.closure_report_template_path = path
            self.closure_report_template_label.config(text=os.path.basename(path))

    def select_zip(self):
        path = filedialog.askopenfilename(filetypes=[("ZIP files", "*.zip")])
        if path:
            self.zip_path = path
            self.zip_label.config(text=os.path.basename(path))

    def generate_excel_report(self):
        if not self.excel_path or not self.zip_path:
            messagebox.showerror("错误", "请先选择隐患整改通知单和图片压缩包！")
            return

        try:
            output_path = os.path.splitext(self.excel_path)[0] + "_带图片.xlsx"
            errors = embed_images_to_excel(self.excel_path, self.zip_path, output_path)

            # 弹出成功消息
            msg = f"处理完成！\n输出文件：\n{output_path}"
            if errors:
                msg += "\n\n⚠️ 部分图片插入失败，详见下方错误信息。"
                messagebox.showinfo("处理完成（含警告）", msg)
                # 弹出错误详情
                error_text = "\n".join(errors[:20])  # 最多显示20条，避免太长
                if len(errors) > 20:
                    error_text += f"\n... 还有 {len(errors) - 20} 条错误未显示"
                messagebox.showwarning("插入失败详情", error_text)
            else:
                messagebox.showinfo("成功", msg)

        except Exception as e:
            messagebox.showerror("严重错误", f"程序运行失败：\n{str(e)}")

    def generate_check_report(self):
        if not self.excel_path or not self.check_report_template_path or not self.zip_path:
            messagebox.showerror("错误", "请先选择隐患整改通知单、检查报告模板和图片压缩包！")
            return

        try:
            output_path = os.path.splitext(self.check_report_template_path)[0] + "_检查报告.docx"
            generate_check_report(self.excel_path, self.check_report_template_path, self.zip_path, output_path)
            messagebox.showinfo("成功", f"检查报告生成完成！\n输出文件：\n{output_path}")

        except Exception as e:
            messagebox.showerror("严重错误", f"程序运行失败：\n{str(e)}")

    def generate_closure_report(self):
        if not self.excel_path or not self.closure_report_template_path or not self.zip_path:
            messagebox.showerror("错误", "请先选择隐患整改通知单、闭环报告模板和图片压缩包！")
            return

        try:
            output_path = os.path.splitext(self.closure_report_template_path)[0] + "_闭环报告.docx"
            generate_closure_report(self.excel_path, self.closure_report_template_path, self.zip_path, output_path)
            messagebox.showinfo("成功", f"闭环报告生成完成！\n输出文件：\n{output_path}")

        except Exception as e:
            messagebox.showerror("严重错误", f"程序运行失败：\n{str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()
